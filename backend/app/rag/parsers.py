from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.exceptions import (
    DocumentProcessingError,
    UnsupportedMediaTypeError,
)


@dataclass(
    frozen=True,
    slots=True,
)
class ParsedSection:
    """
    One semantic section extracted from a document.

    section_path example:
        ("Employee Handbook", "Leave Policy", "Annual Leave")

    content_type:
        prose
        table
    """

    text: str

    page_number: int | None = None

    title: str | None = None

    section_path: tuple[str, ...] = ()

    content_type: str = "prose"

    metadata: dict[str, Any] = field(
        default_factory=dict
    )


@dataclass(
    frozen=True,
    slots=True,
)
class ParsedDocument:
    sections: list[ParsedSection]

    @property
    def text(self) -> str:
        return "\n\n".join(
            section.text
            for section in self.sections
            if section.text.strip()
        )


def parse_document(
    *,
    file_bytes: bytes,
    filename: str,
    extension: str,
) -> ParsedDocument:
    normalized_extension = (
        extension.strip().lower()
    )

    if normalized_extension == ".pdf":
        parsed = _parse_pdf(
            file_bytes
        )

    elif normalized_extension == ".docx":
        parsed = _parse_docx(
            file_bytes
        )

    elif normalized_extension == ".md":
        parsed = _parse_markdown(
            file_bytes
        )

    elif normalized_extension == ".txt":
        parsed = _parse_text(
            file_bytes
        )

    else:
        raise UnsupportedMediaTypeError(
            "Files with the "
            f"{normalized_extension or 'unknown'} "
            "extension are not supported"
        )

    if (
        not parsed.sections
        or not parsed.text.strip()
    ):
        raise DocumentProcessingError(
            "No readable text could be "
            f"extracted from {filename}"
        )

    return parsed


# =========================================================
# PDF
# =========================================================


def _parse_pdf(
    file_bytes: bytes,
) -> ParsedDocument:
    if not file_bytes:
        raise DocumentProcessingError(
            "The PDF file is empty"
        )

    try:
        reader = PdfReader(
            BytesIO(file_bytes)
        )

    except (
        PdfReadError,
        ValueError,
        TypeError,
        OSError,
    ) as exc:
        raise DocumentProcessingError(
            "The PDF file is corrupted "
            "or cannot be read"
        ) from exc

    except Exception as exc:
        raise DocumentProcessingError(
            "Unable to read the PDF file"
        ) from exc

    if reader.is_encrypted:
        try:
            decrypt_result = (
                reader.decrypt("")
            )

        except Exception as exc:
            raise DocumentProcessingError(
                "Encrypted PDF files "
                "requiring a password "
                "are not supported"
            ) from exc

        if decrypt_result == 0:
            raise DocumentProcessingError(
                "Encrypted PDF files "
                "requiring a password "
                "are not supported"
            )

    if not reader.pages:
        raise DocumentProcessingError(
            "The PDF file does not "
            "contain any pages"
        )

    sections: list[
        ParsedSection
    ] = []

    for page_index, page in enumerate(
        reader.pages
    ):
        try:
            extracted = (
                page.extract_text()
            )

        except Exception as exc:
            raise DocumentProcessingError(
                "Unable to extract text from "
                f"PDF page {page_index + 1}"
            ) from exc

        text = (
            extracted.strip()
            if extracted
            else ""
        )

        if not text:
            continue

        page_sections = (
            _split_plain_text_sections(
                text=text,
                page_number=(
                    page_index + 1
                ),
            )
        )

        sections.extend(
            page_sections
        )

    if not sections:
        raise DocumentProcessingError(
            "The PDF does not contain "
            "extractable text"
        )

    return ParsedDocument(
        sections=sections
    )


# =========================================================
# DOCX
# =========================================================


def _parse_docx(
    file_bytes: bytes,
) -> ParsedDocument:
    if not file_bytes:
        raise DocumentProcessingError(
            "The DOCX file is empty"
        )

    try:
        document = DocxDocument(
            BytesIO(file_bytes)
        )

    except (
        PackageNotFoundError,
        ValueError,
        KeyError,
        TypeError,
        OSError,
    ) as exc:
        raise DocumentProcessingError(
            "The DOCX file is corrupted "
            "or cannot be read"
        ) from exc

    except Exception as exc:
        raise DocumentProcessingError(
            "Unable to read the DOCX file"
        ) from exc

    sections: list[
        ParsedSection
    ] = []

    heading_stack: list[str] = []

    paragraph_buffer: list[str] = []

    def flush_prose() -> None:
        nonlocal paragraph_buffer

        text = "\n\n".join(
            value
            for value
            in paragraph_buffer
            if value.strip()
        ).strip()

        paragraph_buffer = []

        if not text:
            return

        sections.append(
            ParsedSection(
                text=text,
                page_number=None,
                title=(
                    heading_stack[-1]
                    if heading_stack
                    else None
                ),
                section_path=tuple(
                    heading_stack
                ),
                content_type="prose",
            )
        )

    for block in _iter_docx_blocks(
        document
    ):
        if isinstance(
            block,
            Paragraph,
        ):
            text = (
                block.text.strip()
            )

            if not text:
                continue

            heading_level = (
                _docx_heading_level(
                    block
                )
            )

            if heading_level is not None:
                flush_prose()

                heading_stack = (
                    _update_heading_path(
                        current=heading_stack,
                        heading=text,
                        level=heading_level,
                    )
                )

                continue

            paragraph_buffer.append(
                text
            )

        elif isinstance(
            block,
            Table,
        ):
            flush_prose()

            table_data = (
                _docx_table_to_text(
                    block
                )
            )

            if table_data is None:
                continue

            (
                table_text,
                headers,
                row_count,
            ) = table_data

            sections.append(
                ParsedSection(
                    text=table_text,
                    page_number=None,
                    title=(
                        heading_stack[-1]
                        if heading_stack
                        else "Table"
                    ),
                    section_path=tuple(
                        heading_stack
                    ),
                    content_type="table",
                    metadata={
                        "table_headers": (
                            headers
                        ),
                        "table_row_count": (
                            row_count
                        ),
                    },
                )
            )

    flush_prose()

    if not sections:
        raise DocumentProcessingError(
            "The DOCX file does not "
            "contain readable text"
        )

    return ParsedDocument(
        sections=sections
    )


def _iter_docx_blocks(
    document: Any,
):
    """
    python-docx .paragraphs and .tables lose document order
    when read separately.

    This iterator preserves the original order.
    """

    body = document.element.body

    for child in body.iterchildren():
        tag = child.tag.rsplit(
            "}",
            maxsplit=1,
        )[-1]

        if tag == "p":
            yield Paragraph(
                child,
                document,
            )

        elif tag == "tbl":
            yield Table(
                child,
                document,
            )


def _docx_heading_level(
    paragraph: Paragraph,
) -> int | None:
    style = paragraph.style

    if style is None:
        return None

    style_name = (
        style.name
        or ""
    ).strip()

    match = re.match(
        r"^Heading\s+(\d+)$",
        style_name,
        flags=re.IGNORECASE,
    )

    if match:
        return max(
            1,
            min(
                int(
                    match.group(1)
                ),
                9,
            ),
        )

    if style_name.casefold() == "title":
        return 1

    return None


def _docx_table_to_text(
    table: Table,
) -> tuple[
    str,
    list[str],
    int,
] | None:
    raw_rows: list[
        list[str]
    ] = []

    for row in table.rows:
        cells = [
            _normalize_inline_text(
                cell.text
            )
            for cell in row.cells
        ]

        if any(cells):
            raw_rows.append(
                cells
            )

    if not raw_rows:
        return None

    width = max(
        len(row)
        for row in raw_rows
    )

    first_row = (
        raw_rows[0]
        + [""] * (
            width
            - len(raw_rows[0])
        )
    )

    headers: list[str] = []

    for index, header in enumerate(
        first_row
    ):
        normalized = (
            header.strip()
            or f"Column {index + 1}"
        )

        headers.append(
            normalized
        )

    data_rows = (
        raw_rows[1:]
        if len(raw_rows) > 1
        else raw_rows
    )

    lines = [
        "Table columns: "
        + " | ".join(
            headers
        )
    ]

    for row_index, row in enumerate(
        data_rows,
        start=1,
    ):
        padded = (
            row
            + [""] * (
                width
                - len(row)
            )
        )

        values: list[str] = []

        for header, value in zip(
            headers,
            padded,
            strict=True,
        ):
            cleaned_value = (
                value.strip()
            )

            if not cleaned_value:
                continue

            values.append(
                f"{header}: "
                f"{cleaned_value}"
            )

        if values:
            lines.append(
                f"Row {row_index}: "
                + " | ".join(
                    values
                )
            )

    return (
        "\n".join(lines),
        headers,
        len(data_rows),
    )


# =========================================================
# Markdown
# =========================================================


def _parse_markdown(
    file_bytes: bytes,
) -> ParsedDocument:
    decoded = _decode_text(
        file_bytes
    )

    if not decoded.strip():
        raise DocumentProcessingError(
            "The Markdown file does not "
            "contain readable text"
        )

    lines = decoded.splitlines()

    sections: list[
        ParsedSection
    ] = []

    heading_stack: list[str] = []

    prose_buffer: list[str] = []

    def flush_prose() -> None:
        nonlocal prose_buffer

        text = "\n".join(
            prose_buffer
        ).strip()

        prose_buffer = []

        if not text:
            return

        sections.append(
            ParsedSection(
                text=text,
                title=(
                    heading_stack[-1]
                    if heading_stack
                    else None
                ),
                section_path=tuple(
                    heading_stack
                ),
                content_type="prose",
            )
        )

    index = 0

    while index < len(lines):
        line = lines[index]

        heading_match = re.match(
            r"^\s{0,3}(#{1,6})\s+(.+?)\s*$",
            line,
        )

        if heading_match:
            flush_prose()

            level = len(
                heading_match.group(1)
            )

            heading = (
                heading_match
                .group(2)
                .strip()
                .strip("#")
                .strip()
            )

            if heading:
                heading_stack = (
                    _update_heading_path(
                        current=heading_stack,
                        heading=heading,
                        level=level,
                    )
                )

            index += 1
            continue

        if _is_markdown_table_start(
            lines,
            index,
        ):
            flush_prose()

            table_lines: list[str] = [
                lines[index],
                lines[index + 1],
            ]

            index += 2

            while (
                index < len(lines)
                and "|" in lines[index]
                and lines[index].strip()
            ):
                table_lines.append(
                    lines[index]
                )

                index += 1

            table_data = (
                _markdown_table_to_text(
                    table_lines
                )
            )

            if table_data is not None:
                (
                    table_text,
                    headers,
                    row_count,
                ) = table_data

                sections.append(
                    ParsedSection(
                        text=table_text,
                        title=(
                            heading_stack[-1]
                            if heading_stack
                            else "Table"
                        ),
                        section_path=tuple(
                            heading_stack
                        ),
                        content_type="table",
                        metadata={
                            "table_headers": (
                                headers
                            ),
                            "table_row_count": (
                                row_count
                            ),
                        },
                    )
                )

            continue

        prose_buffer.append(
            line
        )

        index += 1

    flush_prose()

    if not sections:
        raise DocumentProcessingError(
            "The Markdown file does not "
            "contain readable text"
        )

    return ParsedDocument(
        sections=sections
    )


def _is_markdown_table_start(
    lines: list[str],
    index: int,
) -> bool:
    if index + 1 >= len(lines):
        return False

    first = (
        lines[index].strip()
    )

    second = (
        lines[index + 1].strip()
    )

    if "|" not in first:
        return False

    return bool(
        re.match(
            (
                r"^\|?\s*"
                r":?-{3,}:?\s*"
                r"(?:\|\s*:?-{3,}:?\s*)+"
                r"\|?$"
            ),
            second,
        )
    )


def _markdown_table_to_text(
    lines: list[str],
) -> tuple[
    str,
    list[str],
    int,
] | None:
    if len(lines) < 2:
        return None

    def parse_row(
        value: str,
    ) -> list[str]:
        value = (
            value.strip()
            .strip("|")
        )

        return [
            cell.strip()
            for cell
            in value.split("|")
        ]

    headers = parse_row(
        lines[0]
    )

    if not headers:
        return None

    normalized_headers = [
        header
        or f"Column {index + 1}"
        for index, header
        in enumerate(headers)
    ]

    data_rows = [
        parse_row(line)
        for line
        in lines[2:]
        if line.strip()
    ]

    output = [
        "Table columns: "
        + " | ".join(
            normalized_headers
        )
    ]

    for row_index, row in enumerate(
        data_rows,
        start=1,
    ):
        padded = (
            row
            + [""] * (
                len(normalized_headers)
                - len(row)
            )
        )

        values: list[str] = []

        for header, value in zip(
            normalized_headers,
            padded[
                :len(
                    normalized_headers
                )
            ],
            strict=True,
        ):
            if not value.strip():
                continue

            values.append(
                f"{header}: "
                f"{value.strip()}"
            )

        if values:
            output.append(
                f"Row {row_index}: "
                + " | ".join(
                    values
                )
            )

    return (
        "\n".join(output),
        normalized_headers,
        len(data_rows),
    )


# =========================================================
# Plain TXT
# =========================================================


def _parse_text(
    file_bytes: bytes,
) -> ParsedDocument:
    if not file_bytes:
        raise DocumentProcessingError(
            "The text file is empty"
        )

    decoded = _decode_text(
        file_bytes
    )

    normalized = (
        decoded.strip()
    )

    if not normalized:
        raise DocumentProcessingError(
            "The text file does not "
            "contain readable text"
        )

    sections = (
        _split_plain_text_sections(
            text=normalized,
            page_number=None,
        )
    )

    return ParsedDocument(
        sections=sections
    )


# =========================================================
# Plain-text hierarchy detection
# =========================================================


def _split_plain_text_sections(
    *,
    text: str,
    page_number: int | None,
) -> list[ParsedSection]:
    lines = [
        line.strip()
        for line
        in text.splitlines()
    ]

    sections: list[
        ParsedSection
    ] = []

    heading_stack: list[str] = []

    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer

        section_text = "\n".join(
            buffer
        ).strip()

        buffer = []

        if not section_text:
            return

        sections.append(
            ParsedSection(
                text=section_text,
                page_number=(
                    page_number
                ),
                title=(
                    heading_stack[-1]
                    if heading_stack
                    else None
                ),
                section_path=tuple(
                    heading_stack
                ),
                content_type="prose",
            )
        )

    for line in lines:
        if not line:
            if (
                buffer
                and buffer[-1] != ""
            ):
                buffer.append(
                    ""
                )

            continue

        heading_info = (
            _detect_plain_heading(
                line
            )
        )

        if heading_info is not None:
            flush()

            level, heading = (
                heading_info
            )

            heading_stack = (
                _update_heading_path(
                    current=heading_stack,
                    heading=heading,
                    level=level,
                )
            )

            continue

        buffer.append(
            line
        )

    flush()

    if not sections and text.strip():
        sections.append(
            ParsedSection(
                text=text.strip(),
                page_number=(
                    page_number
                ),
            )
        )

    return sections


def _detect_plain_heading(
    line: str,
) -> tuple[
    int,
    str,
] | None:
    value = (
        line.strip()
    )

    if not value:
        return None

    if len(value) > 120:
        return None

    words = value.split()

    if len(words) > 14:
        return None

    numbered = re.match(
        r"^(\d+(?:\.\d+){0,5})[\.)]?\s+(.+)$",
        value,
    )

    if numbered:
        number = (
            numbered.group(1)
        )

        heading = (
            numbered.group(2)
            .strip()
        )

        level = min(
            number.count(".") + 1,
            6,
        )

        return (
            level,
            heading,
        )

    if value.endswith(
        (
            ".",
            "?",
            "!",
            ",",
            ";",
        )
    ):
        return None

    letters = "".join(
        character
        for character in value
        if character.isalpha()
    )

    if (
        letters
        and letters.isupper()
        and len(words) <= 12
    ):
        return (
            1,
            value,
        )

    title_like = all(
        (
            word[:1].isupper()
            or word.casefold()
            in {
                "and",
                "or",
                "of",
                "the",
                "to",
                "for",
                "in",
                "on",
                "a",
                "an",
            }
        )
        for word in words
        if word
    )

    if (
        title_like
        and 1 <= len(words) <= 8
    ):
        return (
            2,
            value,
        )

    return None


# =========================================================
# Helpers
# =========================================================


def _update_heading_path(
    *,
    current: list[str],
    heading: str,
    level: int,
) -> list[str]:
    level = max(
        1,
        level,
    )

    path = list(
        current[
            :level - 1
        ]
    )

    while len(path) < (
        level - 1
    ):
        path.append(
            ""
        )

    path.append(
        heading.strip()
    )

    return [
        value
        for value in path
        if value
    ]


def _normalize_inline_text(
    value: str,
) -> str:
    return (
        " ".join(
            value.split()
        )
    )


def _decode_text(
    file_bytes: bytes,
) -> str:
    if not file_bytes:
        raise DocumentProcessingError(
            "The text file is empty"
        )

    encodings = (
        "utf-8-sig",
        "utf-8",
        "utf-16",
    )

    for encoding in encodings:
        try:
            return file_bytes.decode(
                encoding
            )

        except UnicodeDecodeError:
            continue

    raise DocumentProcessingError(
        "The text file uses an "
        "unsupported character encoding"
    )


def get_extension(
    filename: str,
) -> str:
    return (
        Path(filename)
        .suffix
        .lower()
    )